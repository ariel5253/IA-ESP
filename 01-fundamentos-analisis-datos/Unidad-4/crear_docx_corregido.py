import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Leer el contenido markdown corregido
with open('informe_eda.md', 'r', encoding='utf-8') as f:
    md_content = f.read()

# Crear documento DOCX
doc = Document()

# Configurar estilos
title_style = doc.styles['Title']
heading1_style = doc.styles['Heading 1']
heading2_style = doc.styles['Heading 2']
normal_style = doc.styles['Normal']

# Procesar línea por línea
lines = md_content.split('\n')
i = 0
while i < len(lines):
    line = lines[i].strip()
    
    if line.startswith('# '):
        # Título principal
        p = doc.add_paragraph(line[2:], style=title_style)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif line.startswith('## '):
        # Encabezado nivel 1
        doc.add_paragraph(line[3:], style=heading1_style)
    elif line.startswith('### '):
        # Encabezado nivel 2
        doc.add_paragraph(line[4:], style=heading2_style)
    elif line.startswith('!['):
        # Imagen
        try:
            # Extraer ruta de imagen y descripción
            parts = line.split('](')
            if len(parts) == 2:
                img_desc = parts[0][2:]
                img_path = parts[1].rstrip(')')
                if os.path.exists(img_path):
                    p = doc.add_paragraph()
                    run = p.runs[0] if p.runs else p.add_run()
                    try:
                        run.add_picture(img_path, width=Inches(5))
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception as e:
                        print(f"Error agregando imagen {img_path}: {e}")
                        doc.add_paragraph(f"[Imagen: {img_desc}]")
        except Exception as e:
            print(f"Error procesando línea de imagen: {e}")
    elif line.startswith('*') and not line.startswith('**'):
        # Descripción de figura
        p = doc.add_paragraph(line[1:])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.italic = True
    elif line.startswith('**') and line.endswith('**'):
        # Texto en negrita
        p = doc.add_paragraph()
        run = p.add_run(line[2:-2])
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif line == '---':
        # Separador - agregar espacio
        doc.add_paragraph()
    elif line and not line.startswith('*'):
        # Texto normal
        doc.add_paragraph(line)
    
    i += 1

# Guardar documento corregido
doc.save('informe_eda_corregido.docx')
print('Documento DOCX corregido creado exitosamente') 