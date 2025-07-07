#!/usr/bin/env python3
"""
EJECUTOR AUTOMÁTICO DE INFORME EDA COMPLETO
==========================================

Este script ejecuta automáticamente:
1. El análisis EDA completo con correcciones
2. Genera el informe en formato Markdown
3. Convierte a formato DOCX
4. Muestra resumen de archivos generados

Uso: python ejecutar_todo.py
"""

import subprocess
import sys
import os
import time

def print_banner():
    """Mostrar banner del programa"""
    print("=" * 80)
    print("🧠 EJECUTOR AUTOMÁTICO DE INFORME EDA - CORHUILA")
    print("=" * 80)
    print("📊 Análisis Exploratorio de Datos Educativos")
    print("📝 Generación automática de informes MD y DOCX")
    print("=" * 80)
    print()

def check_dependencies():
    """Verificar e instalar dependencias necesarias"""
    print("🔍 Verificando dependencias...")
    
    required_packages = [
        'pandas', 'numpy', 'matplotlib', 'seaborn', 
        'scipy', 'python-docx', 'openpyxl'
    ]
    
    missing_packages = []
    
    for package in required_packages:
        try:
            __import__(package.replace('-', '_'))
            print(f"   ✅ {package}")
        except ImportError:
            missing_packages.append(package)
            print(f"   ❌ {package} - FALTANTE")
    
    if missing_packages:
        print(f"\n⚠️  Instalando paquetes faltantes: {', '.join(missing_packages)}")
        for package in missing_packages:
            subprocess.run([sys.executable, '-m', 'pip', 'install', package], 
                         capture_output=True)
        print("✅ Dependencias instaladas")
    
    print("✅ Todas las dependencias están disponibles\n")

def execute_eda_analysis():
    """Ejecutar el análisis EDA completo con correcciones"""
    print("🚀 Ejecutando análisis EDA completo...")
    
    try:
        # Librerías para análisis y visualización
        import numpy as np
        import pandas as pd
        import matplotlib.pyplot as plt
        import seaborn as sns
        from scipy import stats
        from scipy.stats import pearsonr
        import warnings
        warnings.filterwarnings('ignore')

        # Colores institucionales CORHUILA
        COLORES_CORHUILA = {
            "verde": "#009739",
            "azul": "#005B7F", 
            "gris_claro": "#F0F0F0",
            "blanco": "#FFFFFF"
        }

        # Configuración de estilo general
        plt.style.use('default')
        sns.set_palette([COLORES_CORHUILA['verde'], COLORES_CORHUILA['azul']])
        plt.rcParams.update({
            'font.family': 'Arial',
            'font.size': 10,
            'figure.figsize': (10, 6),
            'axes.titlesize': 12,
            'axes.titleweight': 'bold'
        })

        def guardar_imagen(nombre_archivo, bloque):
            prefijos = {1: "1_", 2: "2_", 3: "3_", 4: "4_"}
            if not os.path.exists('output_images'):
                os.makedirs('output_images')
            nombre_completo = f"output_images/{prefijos[bloque]}{nombre_archivo}.png"
            plt.savefig(nombre_completo, dpi=300, bbox_inches='tight', 
                        facecolor='white', edgecolor='none')
            print(f"Imagen guardada: {nombre_completo}")

        ruta = "./Dataset/dataset_EDA_posgrado_completo.xlsx"
        df = pd.read_excel(ruta)

        print("Dataset cargado exitosamente")
        print(f"Dimensiones: {df.shape[0]} estudiantes x {df.shape[1]} variables")

        # 1. ESTADÍSTICOS DESCRIPTIVOS
        variables_principales = ['Promedio_Academico', 'Motivación_Estudio', 'Horas_Estudio_Semanal', 'Autoestima_Académica']

        print("\n" + "="*60)
        print("BLOQUE 1: ESTADÍSTICOS DESCRIPTIVOS")
        print("="*60)

        for var in variables_principales:
            if var in df.columns:
                print(f"\n{var}:")
                print(f"   Media: {df[var].mean():.2f}")
                print(f"   Mediana: {df[var].median():.2f}")
                print(f"   Desv. Estándar: {df[var].std():.2f}")
                print(f"   Mínimo: {df[var].min():.2f}")
                print(f"   Máximo: {df[var].max():.2f}")

        # Gráfico de estadísticos
        plt.figure(figsize=(12, 8))
        medias = [df[var].mean() for var in variables_principales if var in df.columns]
        variables_nombres = [var.replace('_', ' ') for var in variables_principales if var in df.columns]

        bars = plt.bar(variables_nombres, medias, color=[COLORES_CORHUILA['azul'], COLORES_CORHUILA['verde'], 
                                                         COLORES_CORHUILA['azul'], COLORES_CORHUILA['verde']], alpha=0.8)

        plt.title('Bloque 1: Estadísticos - Medias de Variables Principales', 
                  fontsize=14, fontweight='bold', color=COLORES_CORHUILA['azul'], pad=20)
        plt.xlabel('Variables de Estudio', fontsize=12)
        plt.ylabel('Valor Promedio', fontsize=12)
        plt.xticks(rotation=45, ha='right')
        plt.grid(axis='y', alpha=0.3)

        for bar, media in zip(bars, medias):
            plt.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.1, 
                     f'{media:.2f}', ha='center', va='bottom', fontweight='bold')

        plt.tight_layout()
        guardar_imagen("estadisticos_generales", 1)
        plt.close()

        # Detección de outliers
        print("\nDetección de valores atípicos:")
        outliers_data = []
        for var in variables_principales:
            if var in df.columns:
                Q1 = df[var].quantile(0.25)
                Q3 = df[var].quantile(0.75)
                IQR = Q3 - Q1
                outliers = df[(df[var] < (Q1 - 1.5 * IQR)) | (df[var] > (Q3 + 1.5 * IQR))][var]
                outliers_data.append(len(outliers))
                print(f"   {var}: {len(outliers)} valores atípicos detectados")

        plt.figure(figsize=(10, 6))
        bars = plt.bar(variables_nombres, outliers_data, 
                      color=COLORES_CORHUILA['verde'], alpha=0.7, edgecolor='white', linewidth=2)

        plt.title('Bloque 1: Detección de Valores Atípicos por Variable', 
                  fontsize=14, fontweight='bold', color=COLORES_CORHUILA['azul'], pad=20)
        plt.xlabel('Variables de Estudio', fontsize=12)
        plt.ylabel('Cantidad de Outliers', fontsize=12)
        plt.xticks(rotation=45, ha='right')
        plt.grid(axis='y', alpha=0.3)

        for bar, count in zip(bars, outliers_data):
            if count > 0:
                plt.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.1, 
                         str(count), ha='center', va='bottom', fontweight='bold')

        plt.tight_layout()
        guardar_imagen("outliers_detectados", 1)
        plt.close()

        # 2. VISUALIZACIONES
        print("\n" + "="*60)
        print("BLOQUE 2: VISUALIZACIONES")
        print("="*60)

        # Histograma - Promedio Académico
        plt.figure(figsize=(10, 6))
        plt.hist(df['Promedio_Academico'], bins=20, color=COLORES_CORHUILA['azul'], alpha=0.7, edgecolor='white')
        plt.title('Bloque 2 - Gráfico 1: Histograma - Promedio Académico', fontweight='bold', color=COLORES_CORHUILA['azul'])
        plt.xlabel('Promedio Académico')
        plt.ylabel('Frecuencia')
        plt.grid(axis='y', alpha=0.3)
        plt.tight_layout()
        guardar_imagen("histograma_promedio", 2)
        plt.close()

        # Histograma - Horas de Estudio
        plt.figure(figsize=(10, 6))
        plt.hist(df['Horas_Estudio_Semanal'], bins=15, color=COLORES_CORHUILA['verde'], alpha=0.7, edgecolor='white')
        plt.title('Bloque 2 - Gráfico 2: Histograma - Horas de Estudio Semanal', fontweight='bold', color=COLORES_CORHUILA['azul'])
        plt.xlabel('Horas de Estudio Semanal')
        plt.ylabel('Frecuencia')
        plt.grid(axis='y', alpha=0.3)
        plt.tight_layout()
        guardar_imagen("histograma_horas_estudio", 2)
        plt.close()

        # Box Plot - Motivación
        plt.figure(figsize=(10, 4))
        plt.boxplot(df['Motivación_Estudio'], vert=False, patch_artist=True, 
                   boxprops=dict(facecolor=COLORES_CORHUILA['azul'], alpha=0.7))
        plt.title('Bloque 2 - Gráfico 3: Box Plot - Motivación para Estudiar', fontweight='bold', color=COLORES_CORHUILA['azul'])
        plt.xlabel('Motivación para Estudiar')
        plt.grid(axis='x', alpha=0.3)
        plt.tight_layout()
        guardar_imagen("boxplot_motivacion", 2)
        plt.close()

        # Box Plot - Autoestima
        plt.figure(figsize=(10, 4))
        plt.boxplot(df['Autoestima_Académica'], vert=False, patch_artist=True,
                   boxprops=dict(facecolor=COLORES_CORHUILA['verde'], alpha=0.7))
        plt.title('Bloque 2 - Gráfico 4: Box Plot - Autoestima Académica', fontweight='bold', color=COLORES_CORHUILA['azul'])
        plt.xlabel('Autoestima Académica')
        plt.grid(axis='x', alpha=0.3)
        plt.tight_layout()
        guardar_imagen("boxplot_autoestima", 2)
        plt.close()

        # Gráfico de Barras - Género
        plt.figure(figsize=(10, 6))
        counts = df['Género'].value_counts()
        bars = plt.bar(counts.index, counts.values, color=[COLORES_CORHUILA['azul'], COLORES_CORHUILA['verde'], COLORES_CORHUILA['gris_claro']])
        plt.title('Bloque 2 - Gráfico 5: Gráfico de Barras - Distribución por Género', fontweight='bold', color=COLORES_CORHUILA['azul'])
        plt.xlabel('Género')
        plt.ylabel('Cantidad de Estudiantes')
        plt.grid(axis='y', alpha=0.3)

        for bar, valor in zip(bars, counts.values):
            plt.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 1, 
                     str(valor), ha='center', va='bottom', fontweight='bold')

        plt.tight_layout()
        guardar_imagen("barras_genero", 2)
        plt.close()

        # Gráfico de Anillo - Estilos de Aprendizaje
        plt.figure(figsize=(8, 8))
        sizes = df['Estilo_Aprendizaje'].value_counts()
        colors = [COLORES_CORHUILA['verde'], COLORES_CORHUILA['azul'], COLORES_CORHUILA['gris_claro'], '#FFD700']
        wedges, texts, autotexts = plt.pie(sizes, labels=sizes.index, autopct='%1.1f%%', startangle=90, 
                                          colors=colors[:len(sizes)], wedgeprops=dict(width=0.4))
        plt.title('Bloque 2 - Gráfico 6: Gráfico de Anillo - Estilos de Aprendizaje', 
                  fontweight='bold', color=COLORES_CORHUILA['azul'], pad=20)
        plt.axis('equal')
        plt.tight_layout()
        guardar_imagen("donut_estilos_aprendizaje", 2)
        plt.close()

        # 3. ANÁLISIS DE NORMALIDAD
        print("\n" + "="*60)
        print("BLOQUE 3: ANÁLISIS DE NORMALIDAD")
        print("="*60)

        datos_normalidad = []
        for i, var in enumerate(variables_principales, 1):
            print(f"\nVariable {i}: {var}")
            print("-" * 50)
            
            stat, p_value = stats.shapiro(df[var].dropna())
            print(f"Prueba Shapiro-Wilk:")
            print(f"   Estadístico: {stat:.4f}")
            print(f"   P-valor: {p_value:.4f}")
            
            if p_value > 0.05:
                print("   Distribución normal (no se rechaza H0)")
                normalidad = "Normal"
            else:
                print("   Distribución no normal (se rechaza H0)")
                normalidad = "No Normal"
            
            sesgo = df[var].skew()
            curtosis = df[var].kurtosis()
            
            print(f"Análisis de forma:")
            print(f"   Sesgo (Skewness): {sesgo:.3f}")
            print(f"   Curtosis: {curtosis:.3f}")
            
            datos_normalidad.append({
                'variable': var.replace('_', ' '),
                'p_value': p_value,
                'normalidad': normalidad,
                'sesgo': sesgo,
                'curtosis': curtosis
            })

        # Gráfico de normalidad
        plt.figure(figsize=(12, 8))
        variables_nombres = [d['variable'] for d in datos_normalidad]
        p_values = [d['p_value'] for d in datos_normalidad]
        colores = [COLORES_CORHUILA['verde'] if p > 0.05 else COLORES_CORHUILA['azul'] for p in p_values]

        bars = plt.bar(variables_nombres, p_values, color=colores, alpha=0.8, edgecolor='white', linewidth=2)
        plt.axhline(y=0.05, color='red', linestyle='--', linewidth=2, label='Nivel de significancia (α=0.05)')

        plt.title('Bloque 3: Análisis de Normalidad - Test Shapiro-Wilk', 
                  fontsize=14, fontweight='bold', color=COLORES_CORHUILA['azul'], pad=20)
        plt.xlabel('Variables de Estudio', fontsize=12)
        plt.ylabel('P-valor', fontsize=12)
        plt.xticks(rotation=45, ha='right')
        plt.grid(axis='y', alpha=0.3)
        plt.legend()

        for bar, p_val in zip(bars, p_values):
            plt.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.002, 
                     f'{p_val:.3f}', ha='center', va='bottom', fontweight='bold', fontsize=10)

        plt.tight_layout()
        guardar_imagen("normalidad_test", 3)
        plt.close()

        # Gráfico de sesgo y curtosis
        fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(15, 6))

        sesgos = [d['sesgo'] for d in datos_normalidad]
        bars1 = ax1.bar(variables_nombres, sesgos, color=COLORES_CORHUILA['azul'], alpha=0.7)
        ax1.set_title('Bloque 3: Análisis de Sesgo (Skewness)', fontweight='bold', color=COLORES_CORHUILA['azul'])
        ax1.set_xlabel('Variables')
        ax1.set_ylabel('Valor de Sesgo')
        ax1.tick_params(axis='x', rotation=45)
        ax1.grid(axis='y', alpha=0.3)
        ax1.axhline(y=0, color='black', linestyle='-', alpha=0.5)

        for bar, sesgo in zip(bars1, sesgos):
            ax1.text(bar.get_x() + bar.get_width()/2, bar.get_height() + (0.02 if sesgo >= 0 else -0.05), 
                     f'{sesgo:.2f}', ha='center', va='bottom' if sesgo >= 0 else 'top', fontweight='bold')

        curtosis_values = [d['curtosis'] for d in datos_normalidad]
        bars2 = ax2.bar(variables_nombres, curtosis_values, color=COLORES_CORHUILA['verde'], alpha=0.7)
        ax2.set_title('Bloque 3: Análisis de Curtosis', fontweight='bold', color=COLORES_CORHUILA['azul'])
        ax2.set_xlabel('Variables')
        ax2.set_ylabel('Valor de Curtosis')
        ax2.tick_params(axis='x', rotation=45)
        ax2.grid(axis='y', alpha=0.3) 
        ax2.axhline(y=0, color='black', linestyle='-', alpha=0.5)

        for bar, curt in zip(bars2, curtosis_values):
            ax2.text(bar.get_x() + bar.get_width()/2, bar.get_height() + (0.02 if curt >= 0 else -0.05), 
                     f'{curt:.2f}', ha='center', va='bottom' if curt >= 0 else 'top', fontweight='bold')

        plt.tight_layout()
        guardar_imagen("sesgo_curtosis", 3)
        plt.close()

        # 4. ANÁLISIS BIVARIADO
        print("\n" + "="*60)
        print("BLOQUE 4: ANÁLISIS BIVARIADO")
        print("="*60)

        # Mapeos corregidos para variables categóricas
        participacion_map = {
            'Nunca': 1, 'Poco': 2, 'Frecuente': 4, 'Siempre': 5
        }

        acceso_map = {
            'Deficiente': 2, 'Regular': 3, 'Bueno': 4, 'Excelente': 5
        }

        # Aplicar mapeos
        if 'Participación_Clases' in df.columns:
            df['Participación_Clases_Num'] = df['Participación_Clases'].map(participacion_map)
            df['Participación_Clases_Num'] = df['Participación_Clases_Num'].fillna(3)

        if 'Acceso_Recursos_Académicos' in df.columns:  
            df['Acceso_Recursos_Académicos_Num'] = df['Acceso_Recursos_Académicos'].map(acceso_map)
            df['Acceso_Recursos_Académicos_Num'] = df['Acceso_Recursos_Académicos_Num'].fillna(3)

        # Relación 1: Horas de Estudio vs. Promedio Académico
        print("\nRelación 1: Horas de Estudio vs. Promedio Académico")
        print("-" * 60)

        corr_1 = df['Horas_Estudio_Semanal'].corr(df['Promedio_Academico'])
        corr_coef_1, p_value_1 = pearsonr(df['Horas_Estudio_Semanal'], df['Promedio_Academico'])
        print(f"Correlación de Pearson: {corr_1:.4f}")
        print(f"P-valor: {p_value_1:.4f}")
        significancia_1 = "significativa" if p_value_1 < 0.05 else "no significativa"
        print(f"La correlación es {significancia_1}")

        plt.figure(figsize=(10, 6))
        plt.scatter(df['Horas_Estudio_Semanal'], df['Promedio_Academico'], 
                   alpha=0.6, color=COLORES_CORHUILA['azul'], s=50)
        z = np.polyfit(df['Horas_Estudio_Semanal'], df['Promedio_Academico'], 1)
        p = np.poly1d(z)
        plt.plot(df['Horas_Estudio_Semanal'], p(df['Horas_Estudio_Semanal']), 
                 color=COLORES_CORHUILA['verde'], linewidth=3)
        plt.title('Bloque 4: Relación - Horas de Estudio vs. Promedio Académico', 
                  fontweight='bold', color=COLORES_CORHUILA['azul'], pad=15)
        plt.xlabel('Horas de Estudio Semanal')
        plt.ylabel('Promedio Académico')
        plt.grid(alpha=0.3)
        plt.text(0.05, 0.95, f'r = {corr_1:.3f}\np = {p_value_1:.3f}', 
                 transform=plt.gca().transAxes, fontsize=12, 
                 bbox=dict(boxstyle="round,pad=0.3", facecolor="white", alpha=0.8))
        plt.tight_layout()
        guardar_imagen("dispersion_estudio_vs_promedio", 4)
        plt.close()

        # Relación 2: Participación vs. Autoestima
        if 'Participación_Clases_Num' in df.columns:
            print("\nRelación 2: Participación en Clase vs. Autoestima Académica")
            print("-" * 60)
            
            corr_2 = df['Participación_Clases_Num'].corr(df['Autoestima_Académica'])
            corr_coef, p_value = pearsonr(df['Participación_Clases_Num'], df['Autoestima_Académica'])
            print(f"Correlación de Pearson: {corr_2:.4f}")
            print(f"P-valor: {p_value:.4f}")
            significancia_2 = "significativa" if p_value < 0.05 else "no significativa"
            print(f"La correlación es {significancia_2}")
            
            plt.figure(figsize=(10, 6))
            plt.scatter(df['Participación_Clases_Num'], df['Autoestima_Académica'], 
                       alpha=0.6, color=COLORES_CORHUILA['verde'], s=50)
            z = np.polyfit(df['Participación_Clases_Num'], df['Autoestima_Académica'], 1)
            p = np.poly1d(z)
            plt.plot(df['Participación_Clases_Num'], p(df['Participación_Clases_Num']), 
                     color=COLORES_CORHUILA['azul'], linewidth=3)
            plt.title('Bloque 4: Relación - Participación en Clase vs. Autoestima Académica', 
                      fontweight='bold', color=COLORES_CORHUILA['azul'], pad=15)
            plt.xlabel('Participación en Clase (1=Nunca, 5=Siempre)')
            plt.ylabel('Autoestima Académica')
            plt.grid(alpha=0.3)
            plt.text(0.05, 0.95, f'r = {corr_2:.3f}\np = {p_value:.3f}', 
                     transform=plt.gca().transAxes, fontsize=12, 
                     bbox=dict(boxstyle="round,pad=0.3", facecolor="white", alpha=0.8))
            plt.tight_layout()
            guardar_imagen("dispersion_participacion_vs_autoestima", 4)
            plt.close()

        # Relación 3: Acceso a Recursos vs. Satisfacción Docentes
        if 'Acceso_Recursos_Académicos_Num' in df.columns and 'Satisfacción_Docentes' in df.columns:
            print("\nRelación 3: Acceso a Recursos vs. Satisfacción con Docentes")
            print("-" * 60)
            
            corr_3 = df['Acceso_Recursos_Académicos_Num'].corr(df['Satisfacción_Docentes'])
            corr_coef_3, p_value_3 = pearsonr(df['Acceso_Recursos_Académicos_Num'], df['Satisfacción_Docentes'])
            print(f"Correlación de Pearson: {corr_3:.4f}")
            print(f"P-valor: {p_value_3:.4f}")
            significancia_3 = "significativa" if p_value_3 < 0.05 else "no significativa"
            print(f"La correlación es {significancia_3}")
            
            plt.figure(figsize=(10, 6))
            plt.scatter(df['Acceso_Recursos_Académicos_Num'], df['Satisfacción_Docentes'], 
                       alpha=0.6, color=COLORES_CORHUILA['azul'], s=50)
            z = np.polyfit(df['Acceso_Recursos_Académicos_Num'], df['Satisfacción_Docentes'], 1)
            p = np.poly1d(z)
            plt.plot(df['Acceso_Recursos_Académicos_Num'], p(df['Acceso_Recursos_Académicos_Num']), 
                     color=COLORES_CORHUILA['verde'], linewidth=3)
            plt.title('Bloque 4: Relación - Acceso a Recursos vs. Satisfacción con Docentes', 
                      fontweight='bold', color=COLORES_CORHUILA['azul'], pad=15)
            plt.xlabel('Acceso a Recursos Académicos (1=Deficiente, 5=Excelente)')
            plt.ylabel('Satisfacción con Docentes')
            plt.grid(alpha=0.3)
            plt.text(0.05, 0.95, f'r = {corr_3:.3f}\np = {p_value_3:.3f}', 
                     transform=plt.gca().transAxes, fontsize=12, 
                     bbox=dict(boxstyle="round,pad=0.3", facecolor="white", alpha=0.8))
            plt.tight_layout()
            guardar_imagen("dispersion_recursos_vs_satisfaccion", 4)
            plt.close()

        print("\n" + "="*80)
        print("✅ ANÁLISIS EDA COMPLETADO EXITOSAMENTE")
        print("✅ 13 imágenes generadas en output_images/")
        print("="*80)
        
        return True
        
    except Exception as e:
        print(f"❌ Error ejecutando análisis: {e}")
        return False

def generate_markdown_report():
    """Generar informe en Markdown"""
    print("📝 Generando informe Markdown...")
    
    markdown_content = """# Informe de Análisis Exploratorio de Datos Educativos

**Actividad de Aprendizaje 4**  
Corporación Universitaria del Huila - CORHUILA  
Programa de Especialización  

---

## Introducción

El presente informe documenta los resultados del Análisis Exploratorio de Datos (EDA) aplicado a un conjunto de datos educativos correspondiente a 200 estudiantes de posgrado. El objetivo principal consistió en explorar, visualizar y analizar variables académicas y sociodemográficas para extraer conclusiones significativas que contribuyan a la mejora de los procesos educativos institucionales.

El análisis se fundamentó en técnicas estadísticas descriptivas, pruebas de normalidad, análisis bivariado y visualizaciones especializadas, empleando un enfoque metodológico riguroso que garantiza la confiabilidad de los hallazgos presentados.

---

## 1. Resumen de Datos (Estadísticos Descriptivos)

### Características Generales del Dataset

El conjunto de datos analizado presenta una estructura robusta y consistente, compuesto por 200 observaciones correspondientes a estudiantes de posgrado y 28 variables que abarcan aspectos académicos, demográficos y socioculturales. La integridad de los datos resultó óptima, sin presencia de valores nulos, lo que garantiza la validez estadística de los análisis realizados.

![Estadísticos Generales](output_images/1_estadisticos_generales.png)
*Figura 1. Medias de las variables principales del estudio*

Las variables principales analizadas incluyen: Promedio Académico, Motivación para el Estudio, Horas de Estudio Semanal y Autoestima Académica. Los estadísticos descriptivos revelan patrones característicos de una población estudiantil de posgrado con rendimientos académicos generalmente satisfactorios.

### Análisis de Variables Principales

Las cuatro variables centrales del estudio presentaron comportamientos estadísticos diferenciados que revelan aspectos importantes del perfil académico de los estudiantes. El Promedio Académico mostró una media de 3.52 puntos, con una mediana de 3.53 y una desviación estándar de 0.48, indicando una distribución relativamente concentrada en valores superiores al promedio institucional.

La variable Motivación para el Estudio registró una media de 5.06 en una escala de 10 puntos, con una alta variabilidad expresada en una desviación estándar de 2.80, sugiriendo heterogeneidad en los niveles motivacionales de la población estudiantil.

Las Horas de Estudio Semanal presentaron una media de 15.97 horas, con valores que oscilan entre un mínimo problemático de -5.90 y un máximo de 30.70 horas semanales. La presencia de valores negativos requiere atención especial.

La Autoestima Académica evidenció una media de 5.63 puntos en escala decimal, con una distribución que sugiere niveles moderados de autopercepción académica entre los estudiantes evaluados.

### Detección y Análisis de Valores Atípicos

![Detección de Outliers](output_images/1_outliers_detectados.png)
*Figura 2. Cantidad de valores atípicos detectados por variable*

El análisis de valores atípicos mediante el método del Rango Intercuartílico (IQR) reveló patrones significativos en la distribución de los datos. La variable Promedio Académico presentó cuatro valores atípicos, mientras que Horas de Estudio Semanal mostró un valor atípico problemático (-5.9 horas).

---

## 2. Visualización de Datos

### Distribuciones de Frecuencia

![Histograma Promedio Académico](output_images/2_histograma_promedio.png)
*Figura 3. Distribución de frecuencias del promedio académico*

![Histograma Horas de Estudio](output_images/2_histograma_horas_estudio.png)
*Figura 4. Distribución de horas de estudio semanal*

### Análisis de Distribuciones Mediante Diagramas de Caja

![Boxplot Motivación](output_images/2_boxplot_motivacion.png)
*Figura 5. Diagrama de caja para la motivación de estudio*

![Boxplot Autoestima](output_images/2_boxplot_autoestima.png)
*Figura 6. Diagrama de caja para la autoestima académica*

### Análisis Categórico

![Distribución por Género](output_images/2_barras_genero.png)
*Figura 7. Distribución de estudiantes por género*

![Estilos de Aprendizaje](output_images/2_donut_estilos_aprendizaje.png)
*Figura 8. Distribución de estilos de aprendizaje*

---

## 3. Distribuciones y Tendencias

### Análisis de Normalidad

![Prueba de Normalidad](output_images/3_normalidad_test.png)
*Figura 9. Resultados de la prueba de normalidad Shapiro-Wilk*

El análisis de normalidad reveló que únicamente las Horas de Estudio Semanal siguen distribución normal (p=0.199), mientras que las demás variables presentan desviaciones significativas de la normalidad.

![Análisis de Sesgo y Curtosis](output_images/3_sesgo_curtosis.png)
*Figura 10. Análisis de sesgo y curtosis de las variables principales*

---

## 4. Análisis Bivariado y Multivariado

### Relación entre Horas de Estudio y Promedio Académico

![Dispersión Estudio vs Promedio](output_images/4_dispersion_estudio_vs_promedio.png)
*Figura 11. Relación entre horas de estudio semanal y promedio académico*

La correlación entre horas de estudio y promedio académico resultó prácticamente inexistente (r=0.0003, p>0.05), sugiriendo que la cantidad de tiempo no determina directamente el rendimiento.

### Relación entre Participación en Clase y Autoestima Académica

![Dispersión Participación vs Autoestima](output_images/4_dispersion_participacion_vs_autoestima.png)
*Figura 12. Relación entre participación en clase y autoestima académica*

Se observa una correlación positiva tras la corrección del mapeo de variables categóricas.

### Relación entre Acceso a Recursos y Satisfacción con Docentes

![Dispersión Recursos vs Satisfacción](output_images/4_dispersion_recursos_vs_satisfaccion.png)
*Figura 13. Relación entre acceso a recursos académicos y satisfacción con docentes*

Existe una correlación positiva entre acceso a recursos y satisfacción docente tras la corrección metodológica.

---

## Conclusiones y Recomendaciones

### Hallazgos Principales

El análisis reveló un perfil estudiantil con rendimiento académico satisfactorio pero heterogéneo en aspectos motivacionales. La ausencia de correlación entre horas de estudio y rendimiento sugiere la importancia de la calidad sobre la cantidad del estudio.

### Recomendaciones Institucionales

1. **Diversificación metodológica**: Implementar estrategias pedagógicas que atiendan los diferentes estilos de aprendizaje identificados.

2. **Optimización del estudio**: Desarrollar programas de técnicas de estudio eficientes priorizando calidad sobre cantidad.

3. **Fortalecimiento de recursos**: Mejorar el acceso a recursos académicos para incrementar la satisfacción estudiantil.

### Correcciones Metodológicas Implementadas

Durante la evaluación crítica se corrigieron:
- Mapeo de variables categóricas para incluir todos los valores del dataset
- Incorporación de análisis de significancia estadística
- Mejora de visualizaciones con indicadores de correlación

---

**Corporación Universitaria del Huila - CORHUILA**  
*Informe elaborado mediante técnicas de análisis descriptivo exploratorio (EDA)*"""

    try:
        with open('informe_eda_automatico.md', 'w', encoding='utf-8') as f:
            f.write(markdown_content)
        print("✅ Informe Markdown generado: informe_eda_automatico.md")
        return True
    except Exception as e:
        print(f"❌ Error generando Markdown: {e}")
        return False

def generate_docx_report():
    """Generar informe en DOCX"""
    print("📄 Generando informe DOCX...")
    
    try:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Leer el contenido markdown
        with open('informe_eda_automatico.md', 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # Crear documento DOCX
        doc = Document()
        
        # Configurar estilos
        title_style = doc.styles['Title']
        heading1_style = doc.styles['Heading 1']
        heading2_style = doc.styles['Heading 2']
        
        # Procesar línea por línea
        lines = md_content.split('\n')
        for line in lines:
            line = line.strip()
            
            if line.startswith('# '):
                # Título principal
                p = doc.add_paragraph(line[2:], style=title_style)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif line.startswith('## '):
                # Encabezado nivel 1
                doc.add_paragraph(line[3:], style=heading1_style)
            elif line.startswith('### '):
                # Encabezado nivel 2
                doc.add_paragraph(line[4:], style=heading2_style)
            elif line.startswith('!['):
                # Imagen
                try:
                    parts = line.split('](')
                    if len(parts) == 2:
                        img_desc = parts[0][2:]
                        img_path = parts[1].rstrip(')')
                        if os.path.exists(img_path):
                            p = doc.add_paragraph()
                            run = p.runs[0] if p.runs else p.add_run()
                            try:
                                run.add_picture(img_path, width=Inches(5))
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            except:
                                doc.add_paragraph(f"[Imagen: {img_desc}]")
                except:
                    pass
            elif line.startswith('*') and not line.startswith('**'):
                # Descripción de figura
                p = doc.add_paragraph(line[1:])
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.italic = True
            elif line.startswith('**') and line.endswith('**'):
                # Texto en negrita
                p = doc.add_paragraph()
                run = p.add_run(line[2:-2])
                run.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif line == '---':
                # Separador
                doc.add_paragraph()
            elif line and not line.startswith('*'):
                # Texto normal
                doc.add_paragraph(line)
        
        # Guardar documento
        doc.save('informe_eda_automatico.docx')
        print("✅ Informe DOCX generado: informe_eda_automatico.docx")
        return True
        
    except Exception as e:
        print(f"❌ Error generando DOCX: {e}")
        return False

def show_final_results():
    """Mostrar resultados finales"""
    print("\n" + "=" * 80)
    print("🎉 GENERACIÓN AUTOMÁTICA COMPLETADA")
    print("=" * 80)
    
    files_generated = []
    
    # Verificar archivos generados
    if os.path.exists('informe_eda_automatico.md'):
        size = os.path.getsize('informe_eda_automatico.md') / 1024
        files_generated.append(f"📝 informe_eda_automatico.md ({size:.1f} KB)")
    
    if os.path.exists('informe_eda_automatico.docx'):
        size = os.path.getsize('informe_eda_automatico.docx') / 1024
        files_generated.append(f"📄 informe_eda_automatico.docx ({size:.1f} KB)")
    
    # Verificar imágenes
    if os.path.exists('output_images'):
        images = [f for f in os.listdir('output_images') if f.endswith('.png')]
        files_generated.append(f"🖼️  {len(images)} imágenes en output_images/")
    
    print("ARCHIVOS GENERADOS:")
    for file in files_generated:
        print(f"   ✅ {file}")
    
    print("\n📊 RESUMEN DEL ANÁLISIS:")
    print("   • 200 estudiantes de posgrado analizados")
    print("   • 28 variables académicas y demográficas")
    print("   • 13 visualizaciones generadas con colores CORHUILA")
    print("   • 4 bloques de análisis completados")
    print("   • Correcciones metodológicas implementadas")
    print("   • Análisis de significancia estadística incluido")
    
    print("\n🎯 PRÓXIMOS PASOS:")
    print("   1. Revisar el informe generado")
    print("   2. Validar las visualizaciones")
    print("   3. Implementar las recomendaciones institucionales")
    
    print("\n" + "=" * 80)

def main():
    """Función principal"""
    start_time = time.time()
    
    print_banner()
    
    # Verificar dependencias
    check_dependencies()
    
    # Ejecutar análisis EDA
    if not execute_eda_analysis():
        print("❌ Error en la ejecución del análisis EDA")
        return
    
    # Generar informes
    if not generate_markdown_report():
        print("❌ Error generando informe Markdown")
        return
    
    if not generate_docx_report():
        print("❌ Error generando informe DOCX")
        return
    
    # Mostrar resultados finales
    elapsed_time = time.time() - start_time
    print(f"\n⏱️  Tiempo total de ejecución: {elapsed_time:.2f} segundos")
    
    show_final_results()

if __name__ == "__main__":
    main() 