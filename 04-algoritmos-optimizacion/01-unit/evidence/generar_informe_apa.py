"""Generador automático de informe APA (6 páginas aprox.) desde notebook ejecutado.

Uso básico dentro del notebook (después de ejecutar todo):

from generar_informe_apa import generar_informe_apa

generar_informe_apa(
    titulo="Asignación Óptima de Horas de Tutoría Basada en Riesgo Académico",
    autores=["Nombre Autor 1", "Nombre Autor 2"],
    institucion="Corporación Universitaria del Huila - Corhuila",
    curso="Algoritmos de Optimización",
    fecha="2025-08-09"
)

Requisitos: pip install python-docx
El script toma variables ya calculadas en el espacio global del notebook si existen:
 base_total, post_total, reduccion_abs, reduccion_pct, H_TOTAL, punto_rd, dual_df, sens_df.
También intenta leer imágenes en ./output_images/ para insertarlas.
"""
from __future__ import annotations
import os, re, math, json, datetime, textwrap
from pathlib import Path
from typing import List, Optional, Dict, Any, Iterable

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError as e:
    raise SystemExit("Falta dependencia python-docx. Ejecute: pip install python-docx")

ROOT = Path(__file__).parent
IMG_DIR = ROOT / 'output_images'
OUTPUT_DOC = ROOT / 'informe_final.docx'

SECTION_ORDER = [
    ("1. Introducción", "introduccion"),
    ("2. Descripción del escenario", "escenario"),
    ("3. Formulación matemática del problema", "formulacion"),
    ("4. Método de resolución", "metodo"),
    ("5. Implementación en Python", "implementacion"),
    ("6. Resultados y análisis", "resultados"),
    ("7. Conclusiones", "conclusiones"),
]

PLACEHOLDER_TEXT = {
    "introduccion": textwrap.dedent(
        """
        Se aborda un problema de asignación óptima de horas de tutoría para mitigar riesgo académico en programas universitarios. La relevancia del problema reside en optimizar el uso de recursos escasos (horas docentes) para maximizar impacto en permanencia estudiantil y eficiencia institucional. La masificación de la educación superior y la presión por indicadores de retención demandan mecanismos cuantitativos reproducibles. El presente trabajo plantea un marco formal capaz de priorizar intervenciones apoyadas en datos integrando dimensiones de programa, área disciplinar y nivel de riesgo. El objetivo general es minimizar el riesgo académico agregado tras la intervención sujeto a limitaciones presupuestales y de capacidad. Los objetivos específicos incluyen: (i) construir métricas de riesgo comparables, (ii) parametrizar un modelo lineal interpretable, (iii) ejecutar análisis de sensibilidad presupuestal y (iv) derivar precios sombra para soporte a la toma de decisiones estratégica.
        """
    ).strip(),
    "escenario": textwrap.dedent(
        """
        El escenario representa una institución con múltiples programas donde cada cohorte exhibe niveles heterogéneos de riesgo académico resumidos mediante indicadores cuantitativos. Las horas de tutoría disponibles constituyen un recurso limitado que puede redistribuirse entre áreas (por ejemplo competencias ciudadanas, matemáticas, comunicación, etc.). Cada área posee una capacidad máxima de absorción y existe además un presupuesto horario global. La intervención genera una reducción marginal del riesgo proporcional a coeficientes calibrados o hipotéticos de efectividad. Este contexto se alinea con procesos de analítica institucional y aprendizaje automático en la medida en que la asignación óptima de recursos humanos mejora la calidad de datos futuros (al reducir deserción y pérdida de seguimiento) y potencialmente la performance de modelos predictivos dependientes de trayectorias académicas completas.
        """
    ).strip(),
    "formulacion": textwrap.dedent(
        """
        Se formula un problema de Programación Lineal (PL) de minimización. Las variables de decisión continuas X_{a} representan horas asignadas al área o cohorte a, mientras que variables auxiliares Y_{a} capturan el riesgo residual para linealizar la función objetivo. La función objetivo minimiza la suma de riesgos posteriores. Las restricciones incluyen: (1) presupuesto horario total (suma de X_{a} ≤ H_TOTAL), (2) capacidades máximas por área (X_{a} ≤ CAP_AREA[a]), (3) piso de equidad (X_{a} ≥ H_MIN) y (4) linealización del riesgo posterior Y_{a} ≥ R_base[a] − K[a]·X_{a} garantizando no negatividad. El modelo es determinista, lineal, con restricciones y variables continuas. Esta estructura promueve interpretabilidad y asegura solvencia mediante métodos simplex o dual simplex. El carácter aditivo y la ausencia de términos cuadráticos eliminan la necesidad de técnicas más complejas en esta iteración.
        """
    ).strip(),
    "metodo": textwrap.dedent(
        """
        El método de resolución seleccionado es Simplex a través de la interfaz PuLP con backend HiGHS/CBC. La elección se fundamenta en: (i) estructura puramente lineal, (ii) necesidad de obtener precios sombra (derivados del óptimo dual) y (iii) tamaño moderado del problema que hace innecesario un enfoque basado en gradiente continuo general o metaheurísticas. Se descarta gradiente descendente por carecer de explotación estructural de vértices y por la potencial dificultad de satisfacer restricciones con igual precisión sin un manejador dedicado. Se menciona que si en futuras extensiones se introducen costos enteros o funciones piecewise no lineales se evaluaría Branch-and-Bound o reformulaciones lineales adicionales.
        """
    ).strip(),
    "implementacion": textwrap.dedent(
        """
        La implementación en Python integra: pandas para manipulación y agregación de datos, PuLP para la modelación y resolución del PL, y seaborn/matplotlib para visualizaciones descriptivas y comparativas pre/post intervención. Se modulariza la derivación de parámetros (H_TOTAL, CAP_AREA, K, H_MIN) para evitar valores codificados, permitiendo regenerar el modelo ante cambios de la base de datos. Se añade un análisis de sensibilidad que recorre presupuestos alternativos construyendo una tabla de reducción marginal de riesgo, identificando el punto de rendimientos decrecientes donde la ganancia incremental cae bajo un umbral relativo del máximo marginal observado. Asimismo, se extraen valores duales y holguras de restricciones clave para estimar el beneficio marginal de relajar límites de capacidad o presupuesto, información crítica para planeación.
        """
    ).strip(),
    "resultados": textwrap.dedent(
        """
        Los resultados cuantifican la reducción del riesgo agregado y exhiben la redistribución eficiente de horas enfocada en áreas con mayor pendiente de mitigación. Las figuras muestran comparativas de riesgo inicial versus residual y barras de asignación de horas. El análisis de sensibilidad evidencia cómo las reducciones marginales decrecen tras un umbral presupuestal (punto RD), sugiriendo que expansiones adicionales del recurso producirían retornos cada vez menores. Los precios sombra indican qué restricciones limitan el óptimo y priorizan inversiones: un precio sombra positivo para el presupuesto total implica valor en ampliar horas; precios sombra en capacidades de áreas revelan cuellos de botella específicos. La tabla de resultados y el resumen cuantitativo integran métricas de reducción absoluta y porcentual, junto con identificación de restricciones activas.
        """
    ).strip(),
    "conclusiones": textwrap.dedent(
        """
        El enfoque de Programación Lineal proporciona una herramienta transparente para asignar recursos tutoriales maximizando la reducción de riesgo académico. Su simplicidad facilita explicabilidad ante actores institucionales. Limitaciones actuales: coeficientes de efectividad K aproximados, ausencia de incertidumbre explícita y omisión de interacciones no lineales entre áreas. Futuras líneas: (i) calibración empírica de K mediante modelos causales o A/B testing, (ii) incorporación de costos diferenciados y restricciones de enteridad, (iii) extensión estocástica para capturar variabilidad de resultados y (iv) integración directa con pipelines de machine learning para retroalimentar predicciones de deserción. Aun con estas limitaciones el modelo ofrece una base cuantitativa robusta para decisiones de priorización.
        """
    ).strip()
}

def _dynamic_section_augments(g: Dict[str, Any]) -> Dict[str, str]:
    """Construye texto adicional interpretativo usando datos reales si existen."""
    out: Dict[str, str] = {}
    try:
        base_total = g.get('base_total')
        post_total = g.get('post_total')
        reduccion_abs = g.get('reduccion_abs')
        reduccion_pct = g.get('reduccion_pct')
        H_TOTAL = g.get('H_TOTAL')
        punto_rd = g.get('punto_rd')
        binding = g.get('binding_list') or []
        prog_summary = g.get('prog_summary')
        prog_hours = g.get('prog_hours') or g.get('pivot_prog_hours')
        dual_df = g.get('dual_df')
        sens_df = g.get('sens_df')
        model = g.get('model')
        X_vars = g.get('X') if isinstance(g.get('X'), dict) else None

        # Descripción cuantitativa del escenario
        if prog_summary is not None:
            n_programas = prog_summary.shape[0]
        else:
            n_programas = None
        if prog_hours is not None:
            try:
                horas_totales_calc = float(prog_hours.select_dtypes('number').sum().sum())
            except Exception:
                horas_totales_calc = None
        else:
            horas_totales_calc = None
        partes_desc = []
        if n_programas:
            partes_desc.append(f"Se analizaron {n_programas} programas/cohortes con riesgo medido.")
        if H_TOTAL is not None:
            partes_desc.append(f"El presupuesto horario disponible fue H_TOTAL={H_TOTAL} horas.")
        if horas_totales_calc and H_TOTAL:
            partes_desc.append(f"La suma de horas asignadas registradas alcanza {horas_totales_calc:,.2f}, cubriendo el {horas_totales_calc / H_TOTAL * 100:,.1f}% del presupuesto declarado.")
        if partes_desc:
            out['escenario'] = ' '.join(partes_desc)

        # Formulación / complejidad
        partes_form = []
        if X_vars:
            partes_form.append(f"Se definen {len(X_vars)} variables de decisión continuas (horas asignadas por área/cohorte).")
        if model is not None:
            try:
                n_constr = len(getattr(model, 'constraints', {}))
                partes_form.append(f"El modelo contiene {n_constr} restricciones explícitas incluyendo presupuesto, capacidades y linealización.")
            except Exception:
                pass
        if dual_df is not None:
            partes_form.append("Los precios sombra provienen de la solución dual del método Simplex.")
        if partes_form:
            out['formulacion'] = ' '.join(partes_form)

        # Resultados narrativos
        partes_res = []
        if base_total is not None and post_total is not None:
            partes_res.append(f"El riesgo agregado pasa de {base_total:,.2f} a {post_total:,.2f}.")
        if reduccion_abs is not None:
            partes_res.append(f"La reducción absoluta es {reduccion_abs:,.2f} ({(reduccion_pct or 0):,.2f}%).")
        if binding:
            partes_res.append("Restricciones vinculantes: " + ', '.join(binding) + ".")
        if punto_rd is not None:
            partes_res.append(f"El punto de rendimientos decrecientes se identificó en H={punto_rd}.")
        # Top programas por reducción y por horas
        if prog_summary is not None:
            cols_lower = [c.lower() for c in prog_summary.columns]
            # intentar detectar columnas de riesgo base/post y reducción
            reduccion_col = None
            for cand in ['reduccion_abs','reduccion','reduccion_total','delta_riesgo']:
                if cand in cols_lower:
                    reduccion_col = prog_summary.columns[cols_lower.index(cand)]
                    break
            if reduccion_col:
                try:
                    top_red = prog_summary.sort_values(reduccion_col, ascending=False).head(3)
                    lista = [str(r.iloc[0]) for _, r in top_red.iterrows()]
                    partes_res.append("Mayores reducciones observadas en: " + ', '.join(lista) + ".")
                except Exception:
                    pass
        if prog_hours is not None:
            try:
                ph = prog_hours.copy()
                numeric_cols = ph.select_dtypes('number')
                if not numeric_cols.empty:
                    ph['TOTAL_H'] = numeric_cols.sum(axis=1)
                    top_h = ph.sort_values('TOTAL_H', ascending=False).head(3)
                    labels = [str(idx) for idx in top_h.index]
                    partes_res.append("Mayor asignación de horas en: " + ', '.join(labels) + ".")
            except Exception:
                pass
        if dual_df is not None:
            try:
                actives = dual_df[dual_df['holgura'].abs() < 1e-6]
                if not actives.empty:
                    partes_res.append("Precios sombra relevantes (holgura≈0): " + ', '.join(f"{r.restriccion}={r.precio_sombra:.3f}" for _, r in actives.head(4).iterrows()) + ".")
            except Exception:
                pass
        if sens_df is not None and 'marginal' in sens_df.columns:
            try:
                partes_res.append(f"Marginal máxima: {sens_df['marginal'].max():,.4f} – marginal final: {sens_df['marginal'].iloc[-1]:,.4f}.")
            except Exception:
                pass
        if partes_res:
            out['resultados'] = ' '.join(partes_res)

        # Conclusiones adaptativas
        partes_conc = []
        if reduccion_pct is not None:
            partes_conc.append(f"El modelo logra una disminución relativa del {reduccion_pct:,.2f}% del riesgo agregado bajo las restricciones actuales.")
        if binding:
            partes_conc.append("Las restricciones activas sugieren focos de inversión para incrementos marginales de impacto.")
        if dual_df is not None:
            partes_conc.append("Los precios sombra proveen guía cuantitativa para priorizar ampliaciones presupuestales o de capacidad.")
        if partes_conc:
            out['conclusiones'] = ' '.join(partes_conc)
    except Exception:
        pass
    return out

def _clean_html(text: str) -> str:
    if not text:
        return ""
    # Quitar etiquetas HTML básicas
    text = re.sub(r'<br\s*/?>', '\n', text, flags=re.I)
    text = re.sub(r'<[^>]+>', '', text)
    return text.strip()

def _normal_paragraph(doc, text: str, bold: bool=False, italic: bool=False, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    style = doc.styles['Normal']
    font = run.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    if align:
        p.alignment = align
    p.paragraph_format.line_spacing = 2
    p.paragraph_format.space_after = Pt(0)
    return p

def _section_heading(doc, title: str):
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.line_spacing = 2
    p.paragraph_format.space_after = Pt(0)

def _add_code_block(doc, code: str, caption: str | None = None):
    # Insert code as preformatted paragraphs (APA allows monospaced exceptions)
    for line in code.rstrip().splitlines():
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(0)
    if caption:
        _normal_paragraph(doc, caption, italic=True)

def _extract_code_snippets(nb_path: Path, patterns: list[str], max_chars: int = 2200,
                           exclude_patterns: Optional[list[str]] = None) -> list[str]:
    """Extrae fragmentos de celdas de código que contengan cualquiera de los patrones.
    Omite celdas que contengan patrones de exclusión (por ejemplo el bloque grande de construcción del modelo).
    Limita el total de caracteres para mantener extensión manejable.
    """
    try:
        import json
        data = json.loads(nb_path.read_text(encoding='utf-8'))
    except Exception:
        return []
    snippets = []
    total = 0
    for cell in data.get('cells', []):
        if cell.get('cell_type') != 'code':
            continue
        src = ''.join(cell.get('source', ''))
        if exclude_patterns:
            lowered = src.lower()
            if any(pat.lower() in lowered for pat in exclude_patterns):
                continue
        if any(pat in src for pat in patterns):
            cleaned = src.strip()
            if not cleaned:
                continue
            length = len(cleaned)
            if total + length > max_chars:
                remaining = max(0, max_chars - total)
                if remaining < 100:
                    break
                cleaned = cleaned[:remaining] + '\n# ... truncado ...'
                snippets.append(cleaned)
                break
            snippets.append(cleaned)
            total += length
    return snippets


def _metrics_block(globals_dict: Dict[str, Any]) -> List[str]:
    base_total = globals_dict.get('base_total')
    post_total = globals_dict.get('post_total')
    reduccion_abs = globals_dict.get('reduccion_abs')
    reduccion_pct = globals_dict.get('reduccion_pct')
    H_TOTAL = globals_dict.get('H_TOTAL')
    punto_rd = globals_dict.get('punto_rd')
    lines = []
    if base_total is not None and post_total is not None:
        lines.append(f"Riesgo base total: {base_total:,.2f}")
        lines.append(f"Riesgo posterior total: {post_total:,.2f}")
    if reduccion_abs is not None:
        lines.append(f"Reducción absoluta: {reduccion_abs:,.2f}")
    if reduccion_pct is not None:
        lines.append(f"Reducción porcentual: {reduccion_pct:,.2f}%")
    if H_TOTAL is not None:
        lines.append(f"Presupuesto horario (H_TOTAL): {H_TOTAL}")
    if punto_rd is not None:
        lines.append(f"Punto rendimientos decrecientes (H): {punto_rd}")
    # Duales
    dual_df = globals_dict.get('dual_df')
    if dual_df is not None:
        try:
            # Mostrar solo las primeras 4 restricciones
            for idx, row in dual_df.head(4).iterrows():
                lines.append(f"Shadow {row['restriccion']}: π={row['precio_sombra']:.4f} holgura={row['holgura']:.2f}")
        except Exception:
            pass
    return lines


def _insert_metrics(doc, globals_dict: Dict[str, Any]):
    lines = _metrics_block(globals_dict)
    if not lines:
        return
    _section_heading(doc, "Resumen cuantitativo principal")
    for ln in lines:
        _normal_paragraph(doc, ln)


def _gather_image_paths(max_images: int = 12) -> List[Path]:
    candidates: List[Path] = []
    search_dirs = [IMG_DIR, Path.cwd() / 'output_images']
    seen = set()
    for d in search_dirs:
        if d.exists():
            for p in sorted(d.glob('*')):
                if p.suffix.lower() in {'.png', '.jpg', '.jpeg'} and p.name not in seen:
                    candidates.append(p)
                    seen.add(p.name)
    return candidates[:max_images]

def _insert_images(doc, max_images: int = 8):
    images = _gather_image_paths(max_images=max_images)
    if not images:
        return
    _section_heading(doc, "Figuras seleccionadas")
    for i, img in enumerate(images, start=1):
        try:
            # Escalar ancho si es muy grande: Word por defecto ajusta, se deja None.
            doc.add_picture(str(img))
            cap = doc.add_paragraph()
            run = cap.add_run(f"Figura {i}. {img.stem.replace('_',' ')}")
            run.font.name = 'Times New Roman'; run.font.size = Pt(11); run.italic = True
            cap.paragraph_format.line_spacing = 2
        except Exception:
            continue


def _insert_dataframe(doc, df, title: str, max_rows: int = 25, note: str | None = None):
    try:
        import pandas as pd  # noqa
    except Exception:
        return
    if df is None:
        return
    if max_rows and len(df) > max_rows:
        df_show = df.head(max_rows).copy()
        truncated = True
    else:
        df_show = df
        truncated = False
    _section_heading(doc, title)
    table = doc.add_table(rows=1, cols=len(df_show.columns))
    hdr_cells = table.rows[0].cells
    for j, col in enumerate(df_show.columns):
        hdr_cells[j].text = str(col)
    for _, row in df_show.iterrows():
        cells = table.add_row().cells
        for j, col in enumerate(df_show.columns):
            cells[j].text = str(row[col])
    table.style = 'Table Grid'
    if truncated:
        _normal_paragraph(doc, f"(Mostradas {len(df_show)} de {len(df)} filas)", italic=True)
    if note:
        _normal_paragraph(doc, note, italic=True)

def _insert_key_tables(doc, g: Dict[str, Any]):
    # Programa / riesgo
    if g.get('prog_summary') is not None:
        _insert_dataframe(doc, g['prog_summary'], 'Tabla: Resumen de Riesgo por Programa', max_rows=30)
    if g.get('prog_hours') is not None:
        _insert_dataframe(doc, g['prog_hours'], 'Tabla: Horas Asignadas por Programa y Área', max_rows=40)
    elif g.get('pivot_prog_hours') is not None:
        _insert_dataframe(doc, g['pivot_prog_hours'], 'Tabla: Horas Asignadas (Pivot)', max_rows=40)
    if g.get('dual_df') is not None:
        _insert_dataframe(doc, g['dual_df'], 'Tabla: Valores Dual (Precios Sombra)', max_rows=60)
    if g.get('sens_df') is not None:
        _insert_dataframe(doc, g['sens_df'], 'Tabla: Sensibilidad Presupuestal', max_rows=30, note='Incluye reducción marginal estimada.')

def _insert_metrics_table(doc, g: Dict[str, Any]):
    metrics = {
        'Riesgo base total': g.get('base_total'),
        'Riesgo posterior total': g.get('post_total'),
        'Reducción absoluta': g.get('reduccion_abs'),
        'Reducción porcentual (%)': g.get('reduccion_pct'),
        'Presupuesto horario (H_TOTAL)': g.get('H_TOTAL'),
        'Punto rendimientos decrecientes (H)': g.get('punto_rd'),
        'Máx. marginal reducción': g.get('max_marg') or g.get('max_marg_calc'),
        'Marginal actual': g.get('marginal_actual'),
        'Restricciones vinculantes': ', '.join(g.get('binding_list', []) or [])
    }
    # Métricas adicionales derivadas por programa si existen
    prog_summary = g.get('prog_summary')
    if prog_summary is not None:
        try:
            cols_lower = [c.lower() for c in prog_summary.columns]
            # localizar columnas de riesgo base/post
            base_col = next((prog_summary.columns[i] for i,c in enumerate(cols_lower) if 'base' in c and 'riesgo' in c), None)
            post_col = next((prog_summary.columns[i] for i,c in enumerate(cols_lower) if ('post' in c or 'residual' in c) and 'riesgo' in c), None)
            if base_col and post_col:
                diff_series = prog_summary[base_col] - prog_summary[post_col]
                metrics['Promedio reducción (programa)'] = float(diff_series.mean())
                metrics['DesvEst reducción (programa)'] = float(diff_series.std())
                improved = (diff_series > 0).sum()
                totalp = len(diff_series)
                metrics['% Programas con mejora'] = improved / totalp * 100 if totalp else None
        except Exception:
            pass
    # Filtrar None
    metrics = {k: v for k, v in metrics.items() if v is not None}
    if not metrics:
        return
    _section_heading(doc, 'Tabla: Métricas Globales del Modelo')
    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = 'Métrica'
    hdr[1].text = 'Valor'
    for k, v in metrics.items():
        row = table.add_row().cells
        row[0].text = k
        if isinstance(v, float):
            row[1].text = f"{v:,.4f}"
        else:
            row[1].text = str(v)
    table.style = 'Table Grid'


def generar_informe_apa(titulo: str, autores: List[str] | None, institucion: str, curso: str, fecha: Optional[str]=None,
                         fuente_referencias: Optional[List[str]] = None,
                         globals_dict: Optional[Dict[str, Any]] = None,
                         salida: Path | str = OUTPUT_DOC,
                         notebook_path: str | Path = 'tarea1_opt_corhuila.ipynb',
                         docente: Optional[str] = None,
                         include_code_snippets: bool = False,
                         auto_generate_figs: bool = True):
    """Genera informe Word (APA básico) aprovechando variables y resultados del entorno.
    No re-ejecuta el notebook: se basa en datos ya presentes.
    """
    if globals_dict is None:
        # Capturar espacio global del intérprete (si se llama desde notebook)
        globals_dict = globals()

    # Derivar métricas faltantes si es posible
    def _derive_missing(g: Dict[str, Any]):
        try:
            prog_summary = g.get('prog_summary')
            if prog_summary is not None:
                cols_lower = [c.lower() for c in prog_summary.columns]
                # base risk column
                base_idx = next((i for i,c in enumerate(cols_lower) if 'riesgo' in c and 'base' in c), None)
                post_idx = next((i for i,c in enumerate(cols_lower) if ('post' in c or 'residual' in c) and 'riesgo' in c), None)
                if g.get('base_total') is None and base_idx is not None:
                    g['base_total'] = float(prog_summary.iloc[:, base_idx].sum())
                if g.get('post_total') is None and post_idx is not None:
                    g['post_total'] = float(prog_summary.iloc[:, post_idx].sum())
                if g.get('reduccion_abs') is None and g.get('base_total') is not None and g.get('post_total') is not None:
                    g['reduccion_abs'] = g['base_total'] - g['post_total']
                if g.get('reduccion_pct') is None and g.get('base_total') not in (None,0) and g.get('reduccion_abs') is not None:
                    g['reduccion_pct'] = g['reduccion_abs'] / g['base_total'] * 100
            # Sensibilidad
            sens_df = g.get('sens_df')
            if sens_df is not None and 'marginal' in sens_df.columns:
                if g.get('max_marg') is None:
                    try:
                        g['max_marg'] = float(sens_df['marginal'].max())
                    except Exception:
                        pass
                if g.get('marginal_actual') is None:
                    try:
                        g['marginal_actual'] = float(sens_df['marginal'].iloc[-1])
                    except Exception:
                        pass
            # Binding constraints from dual_df
            dual_df = g.get('dual_df')
            if dual_df is not None and not g.get('binding_list'):
                try:
                    g['binding_list'] = [str(r.restriccion) for _, r in dual_df.iterrows() if abs(r.holgura) < 1e-6][:10]
                except Exception:
                    pass
        except Exception:
            pass
    _derive_missing(globals_dict)
    if fecha is None:
        fecha = datetime.date.today().isoformat()

    doc = Document()
    # Configuración de estilos APA básicos
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Portada (APA 7 básica) centrada
    if not autores:
        # Intentar leer de variable global AUTORES
        posibles = globals_dict.get('AUTORES') if globals_dict else None
        if isinstance(posibles, (list, tuple)):
            autores = list(posibles)
        else:
            autores = ['Autor No Definido']
    cover_items = [titulo.upper(), *autores, institucion, curso]
    if docente:
        cover_items.append(docente)
    cover_items.append(fecha)
    for i, line in enumerate(cover_items):
        # Añadir espacio antes para empujar a mitad de página
        if i == 0:
            for _ in range(6):
                _normal_paragraph(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER)
        _normal_paragraph(doc, line, bold=(i==0), align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # Secciones
    nb_path = Path(notebook_path)
    code_snips = []
    if include_code_snippets and nb_path.exists():
        code_snips = _extract_code_snippets(
            nb_path,
            patterns=['LpProblem','objective','constraints','model +=','solve','sens_df','dual_df'],
            exclude_patterns=[
                'Construcción del modelo de Programación Lineal',
                'Construccion del modelo de Programacion Lineal'
            ]
        )

    def _auto_generate_figs(g: Dict[str, Any]):
        # Solo si carpeta vacía
        existing = _gather_image_paths(max_images=2)
        if existing:
            return
        try:
            import matplotlib.pyplot as plt
            import seaborn as sns
            import pandas as pd  # noqa
        except Exception:
            return
        IMG_DIR.mkdir(exist_ok=True)
        # 1. Riesgo base vs post (si prog_summary tiene columnas adecuadas)
        prog_summary = g.get('prog_summary')
        if prog_summary is not None and not prog_summary.empty:
            cols_lower = [c.lower() for c in prog_summary.columns]
            try:
                base_idx = next((i for i,c in enumerate(cols_lower) if 'riesgo' in c and 'base' in c), None)
                post_idx = next((i for i,c in enumerate(cols_lower) if ('post' in c or 'residual' in c) and 'riesgo' in c), None)
                if base_idx is not None and post_idx is not None:
                    df_plot = prog_summary.iloc[:, [base_idx, post_idx]].copy()
                    df_plot.columns = ['Riesgo Base','Riesgo Post']
                    df_plot_mean = df_plot.mean()
                    plt.figure(figsize=(5,3))
                    sns.barplot(x=df_plot_mean.index, y=df_plot_mean.values, palette='viridis')
                    plt.title('Promedio Riesgo Base vs Post')
                    plt.ylabel('Riesgo')
                    plt.tight_layout()
                    plt.savefig(IMG_DIR / 'fig_riesgo_promedio.png', dpi=150)
                    plt.close()
            except Exception:
                pass
        # 2. Sensibilidad presupuestal curva
        sens_df = g.get('sens_df')
        if sens_df is not None and not sens_df.empty and 'H' in sens_df.columns:
            try:
                y_col = None
                for cand in ['reduccion','reduccion_abs','reduccion_total','delta']:
                    if cand in sens_df.columns:
                        y_col = cand
                        break
                if y_col:
                    plt.figure(figsize=(5,3))
                    sns.lineplot(data=sens_df, x='H', y=y_col, marker='o')
                    plt.title('Curva de Reducción vs Presupuesto')
                    plt.xlabel('H')
                    plt.ylabel('Reducción')
                    plt.tight_layout()
                    plt.savefig(IMG_DIR / 'fig_sensibilidad.png', dpi=150)
                    plt.close()
                if 'marginal' in sens_df.columns:
                    plt.figure(figsize=(5,3))
                    sns.lineplot(data=sens_df, x='H', y='marginal', marker='o', color='orange')
                    plt.title('Marginal de Reducción')
                    plt.xlabel('H')
                    plt.ylabel('Marginal')
                    plt.tight_layout()
                    plt.savefig(IMG_DIR / 'fig_marginal.png', dpi=150)
                    plt.close()
            except Exception:
                pass
        # 3. Precios sombra
        dual_df = g.get('dual_df')
        if dual_df is not None and not dual_df.empty:
            try:
                top_dual = dual_df.copy()
                if 'precio_sombra' in top_dual.columns:
                    top_dual = top_dual.sort_values('precio_sombra', ascending=False).head(8)
                    plt.figure(figsize=(6,3))
                    sns.barplot(x='precio_sombra', y='restriccion', data=top_dual, palette='mako')
                    plt.title('Precios Sombra Principales')
                    plt.xlabel('π')
                    plt.ylabel('Restricción')
                    plt.tight_layout()
                    plt.savefig(IMG_DIR / 'fig_precios_sombra.png', dpi=150)
                    plt.close()
            except Exception:
                pass
        # 4. Distribución horas (si prog_hours)
        prog_hours = g.get('prog_hours') or g.get('pivot_prog_hours')
        if prog_hours is not None:
            try:
                ph = prog_hours.copy()
                num = ph.select_dtypes('number')
                if not num.empty:
                    top_rows = num.sum(axis=1).sort_values(ascending=False).head(6).index
                    subset = num.loc[top_rows]
                    plt.figure(figsize=(6,3))
                    sns.heatmap(subset, cmap='YlGnBu', annot=False)
                    plt.title('Horas Asignadas (Top)')
                    plt.tight_layout()
                    plt.savefig(IMG_DIR / 'fig_horas_heatmap.png', dpi=150)
                    plt.close()
            except Exception:
                pass

    if auto_generate_figs:
        _auto_generate_figs(globals_dict)

    def _insert_result_values_for_implementation(doc, g: Dict[str, Any]):
        sens_df = g.get('sens_df')
        dual_df = g.get('dual_df')
        prog_summary = g.get('prog_summary')
        bullets: list[str] = []
        if sens_df is not None and not sens_df.empty:
            try:
                H_min = sens_df['H'].min() if 'H' in sens_df.columns else None
                H_max = sens_df['H'].max() if 'H' in sens_df.columns else None
                max_marg = sens_df['marginal'].max() if 'marginal' in sens_df.columns else None
                last_marg = sens_df['marginal'].iloc[-1] if 'marginal' in sens_df.columns else None
                bullets.append(f"Sensibilidad presupuestal evaluada en {len(sens_df)} puntos de H{f' ({H_min} a {H_max})' if H_min is not None and H_max is not None else ''}.")
                if max_marg is not None and last_marg is not None:
                    bullets.append(f"Marginal máxima de reducción {max_marg:,.4f}; marginal final {last_marg:,.4f}.")
            except Exception:
                pass
        if g.get('punto_rd') is not None:
            bullets.append(f"Punto de rendimientos decrecientes detectado en H={g['punto_rd']}.")
        if g.get('base_total') is not None and g.get('post_total') is not None:
            bullets.append(f"Riesgo total base {g['base_total']:,.2f} → posterior {g['post_total']:,.2f} (reducción {g.get('reduccion_abs', 0):,.2f} = {g.get('reduccion_pct',0):,.2f}%).")
        if prog_summary is not None:
            try:
                bullets.append(f"Programas/cohortes modelados: {prog_summary.shape[0]}.")
            except Exception:
                pass
        if dual_df is not None:
            try:
                actives = dual_df[dual_df['holgura'].abs() < 1e-6]
                if not actives.empty:
                    samples = ', '.join(f"{r.restriccion} (π={r.precio_sombra:.3f})" for _, r in actives.head(5).iterrows())
                    bullets.append("Restricciones activas y precios sombra: " + samples + ".")
            except Exception:
                pass
        if g.get('binding_list'):
            bullets.append("Binding: " + ', '.join(g['binding_list']) + ".")
        if bullets:
            _normal_paragraph(doc, 'Resumen cuantitativo de implementación:', bold=True)
            for b in bullets:
                _normal_paragraph(doc, '• ' + b)

    augments = _dynamic_section_augments(globals_dict)

    for title, key in SECTION_ORDER:
        _section_heading(doc, title)
        body = PLACEHOLDER_TEXT.get(key, '')
        if key in augments:
            # Combinar base + dinámico
            body = body + '\n\n' + augments[key]
        _normal_paragraph(doc, body)
        # Añadir fórmula textual en formulación
        if key == 'formulacion':
            formula_lines = [
                'Función objetivo: Min Σ Y_a',
                'Sujeto a:',
                'Y_a >= R_base[a] - K[a]*X_a   (linealización riesgo)',
                'X_a <= CAP_AREA[a]           (capacidad por área)',
                'sum_a X_a <= H_TOTAL         (presupuesto)',
                'X_a >= H_MIN                 (equidad mínima)',
                'X_a, Y_a >= 0'
            ]
            for fl in formula_lines:
                _normal_paragraph(doc, fl)
        if key == 'resultados':
            _insert_metrics_table(doc, globals_dict)
            _insert_images(doc)
            _insert_key_tables(doc, globals_dict)
        if key == 'implementacion':
            if include_code_snippets and code_snips:
                _normal_paragraph(doc, 'Fragmentos clave de código:', bold=True)
                for sn in code_snips[:4]:
                    _add_code_block(doc, sn)
            else:
                _insert_result_values_for_implementation(doc, globals_dict)
    # (Footer automático removido a petición: no se añade nota en conclusiones)
        doc.add_paragraph()  # espaciado

    # Referencias
    _section_heading(doc, "Referencias")
    if not fuente_referencias:
        fuente_referencias = [
            'PuLP Community (2023). PuLP: Linear Programming in Python.',
            'Python Software Foundation (2024). Python Language Reference.',
            'Hunter, J. D. (2007). Matplotlib: A 2D graphics environment. Computing in Science & Engineering.'
        ]
    for ref in fuente_referencias:
        _normal_paragraph(doc, ref)

    # Guardado robusto: si el archivo está abierto en Word, crear nombre alterno
    final_path = Path(salida)
    try:
        doc.save(str(final_path))
    except PermissionError:
        ts = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
        alt = final_path.parent / f"{final_path.stem}_{ts}{final_path.suffix}"
        doc.save(str(alt))
        final_path = alt
    return final_path

if __name__ == '__main__':
    generar_informe_apa(
        titulo="Asignación Óptima de Horas de Tutoría Basada en Riesgo Académico",
        autores=["Autor 1","Autor 2"],
        institucion="Corporación Universitaria del Huila - Corhuila",
        curso="Algoritmos de Optimización"
    )
    print(f"Informe generado en {OUTPUT_DOC}")
